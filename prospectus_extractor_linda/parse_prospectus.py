from __future__ import annotations

from pathlib import Path
from typing import List, Dict, Any

import json
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from config_prospectus import MAX_CHUNK_WORDS


# ─────────────────────────────────────────────────────────────
# 1. Parcours bloc par bloc (paragraphes + tables)
# ─────────────────────────────────────────────────────────────

def iter_block_items(doc: Document):
    """
    Itère sur les éléments de haut niveau du document dans l'ordre :
    Paragraph, puis Table, etc.

    Permet de garder l'ordre réel du texte (contrairement à doc.paragraphs
    et doc.tables séparément).
    """
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def table_to_text(table: Table) -> str:
    """
    Convertit un tableau Word en texte brut lisible.

    Exemple de sortie :

    TABLE:
    Row: Base Currency | ETF Classes | Non-ETF Classes
    Row:               | USD         | USD
    ...
    """
    rows_text = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            # Nettoyage basique
            cell_text = cell.text.replace("\n", " ").strip()
            cells.append(cell_text)
        row_str = " | ".join(cells)
        if row_str.strip():
            rows_text.append(f"Row: {row_str}")

    if not rows_text:
        return ""

    return "TABLE:\n" + "\n".join(rows_text)


# ─────────────────────────────────────────────────────────────
# 2. Détection des sections + accumulation texte + tables
# ─────────────────────────────────────────────────────────────

def is_heading(paragraph: Paragraph) -> bool:
    """
    Heuristique pour détecter un titre de section.

    - Style Word de type 'Heading X'
    - OU texte en MAJUSCULES/assez court (INTRODUCTION, ANNEX, etc.)
    """
    text = paragraph.text.strip()
    if not text:
        return False

    style_name = (paragraph.style.name or "").lower()

    if style_name.startswith("heading"):
        return True

    # Heuristique "titre en majuscules"
    if len(text) < 80 and text.upper() == text:
        return True

    return False


def classify_section(title: str) -> str:
    """
    Classe la section de manière grossière pour aider l'agent.
    """
    t = title.lower()

    if "objective" in t or "policies" in t:
        return "objective"
    if "risk" in t:
        return "risk_profile"
    if "dividend policy" in t or "distribution" in t:
        return "distribution_policy"
    if "charges" in t or "expenses" in t or "fees" in t:
        return "fees"
    if "restriction" in t:
        return "investment_restrictions"
    if "available shares" in t or "registration" in t or "share dealing" in t:
        return "eligibility"
    if "annex" in t or "sustainable investment" in t:
        return "objective"  # partie SFDR / ESG (objective ESG)
    return "other"


def extract_sections_from_prospectus(prospectus_path: str | Path) -> List[Dict[str, Any]]:
    """
    Lit le prospectus Word et retourne une liste de sections :

    [
      {
        "title": "...",
        "text": "texte concaténé (paragraphes + tables)",
      },
      ...
    ]
    """
    doc = Document(str(prospectus_path))

    sections: List[Dict[str, Any]] = []
    current_title = "INTRODUCTION"
    current_buffer: List[str] = []

    def flush_section():
        nonlocal current_title, current_buffer, sections
        if not current_buffer:
            return
        full_text = " ".join(current_buffer).strip()
        if not full_text:
            return
        sections.append({
            "title": current_title.strip(),
            "text": full_text,
        })
        current_buffer = []

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue

            if is_heading(block):
                # On commence une nouvelle section : on flush l'ancienne
                flush_section()
                current_title = text
            else:
                current_buffer.append(text)

        elif isinstance(block, Table):
            # On convertit le tableau en texte et on l'ajoute à la section courante
            table_text = table_to_text(block)
            if table_text.strip():
                current_buffer.append(table_text)

    # Dernière section
    flush_section()

    return sections


# ─────────────────────────────────────────────────────────────
# 3. Chunking des sections (200–500 mots) + métadonnées
# ─────────────────────────────────────────────────────────────

def chunk_section(
    title: str,
    text: str,
    fund_name: str,
    max_words: int = MAX_CHUNK_WORDS,
) -> List[Dict[str, Any]]:
    """
    Découpe une section en sous-chunks de max_words.
    """
    words = text.split()
    chunks: List[Dict[str, Any]] = []

    if not words:
        return chunks

    part_index = 1
    for start in range(0, len(words), max_words):
        piece = " ".join(words[start:start + max_words])
        section_type = classify_section(title)

        chunks.append({
            "section_name": (
                f"{title} (part {part_index})" if len(words) > max_words else title
            ),
            "content": piece,
            "metadata": {
                "fund_name": fund_name,
                "section_type": section_type,
                "original_title": title,
                "part_index": part_index,
            },
        })
        part_index += 1

    return chunks


def parse_prospectus_to_chunks(
    prospectus_path: str | Path,
    fund_name: str,
    max_chunk_words: int = MAX_CHUNK_WORDS,
) -> List[Dict[str, Any]]:
    sections = extract_sections_from_prospectus(prospectus_path)
    all_chunks: List[Dict[str, Any]] = []

    for section in sections:
        title = section["title"]
        text = section["text"]
        section_chunks = chunk_section(
            title=title,
            text=text,
            fund_name=fund_name,
            max_words=max_chunk_words,
        )
        all_chunks.extend(section_chunks)

    return all_chunks


def save_prospectus_chunks_to_json(
    prospectus_path: str | Path,
    fund_name: str,
    output_json: str | Path,
    max_chunk_words: int = MAX_CHUNK_WORDS,
) -> List[Dict[str, Any]]:
    chunks = parse_prospectus_to_chunks(
        prospectus_path=prospectus_path,
        fund_name=fund_name,
        max_chunk_words=max_chunk_words,
    )

    output_json = Path(output_json)
    output_json.parent.mkdir(parents=True, exist_ok=True)

    with output_json.open("w", encoding="utf-8") as f:
        json.dump(chunks, f, ensure_ascii=False, indent=2)

    print(f"[ETAPE 3B] Sauvegarde de {len(chunks)} chunks dans {output_json}")
    return chunks


if __name__ == "__main__":
    # Petit test manuel
    here = Path(__file__).parent
    PROSPECTUS_PATH = here / "prospectus.docx"
    FUND_NAME = "ODDO BHF US Equity Active UCITS ETF"
    OUTPUT_JSON = here / "outputs" / "parsed" / "prospectus_parsed.json"

    save_prospectus_chunks_to_json(
        prospectus_path=PROSPECTUS_PATH,
        fund_name=FUND_NAME,
        output_json=OUTPUT_JSON,
        max_chunk_words=450,
    )
